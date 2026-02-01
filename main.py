# main.py
import os
from docx import Document
from docx.shared import Inches, Pt

from tkinter import filedialog
import tkinter as tk
from tkinter import ttk, messagebox
from tkcalendar import DateEntry
from datetime import date, datetime, timedelta
import sys
from db import *
from create_db import create_database
from tksheet import Sheet
from timedelta_work import timedelta_to_string


# Сначала создаем БД если нужно
create_database()

page_names = ["Главная", "Запись", "График работы", "Услуги", "Клиенты", "Мастера", "Отчеты"]
pages = {}
menu_buttons = {}
current_page = None

# Глобальные переменные для хранения состояния
selected_time_slot = None  # Выбранный временной слот на странице записи
selected_master_id_on_booking_page = None  # Выбранный мастер на странице записи

root = tk.Tk()
root.title("Система управления салоном маникюра")
root.state('zoomed')
root.configure(bg="#F0F0F0")


def quit_app():
    """Выход из приложения"""
    root.quit()
    root.destroy()


def show_page(page_name):
    """Переключение между страницами"""
    global current_page
    if current_page and current_page in menu_buttons:
        menu_buttons[current_page].configure(style="Menu.TButton")
    pages[page_name].tkraise()
    current_page = page_name
    menu_buttons[page_name].configure(style="Selected.Menu.TButton")

    # Обновляем данные при переходе на страницу
    if page_name == "Клиенты":
        refresh_clients_page()
    elif page_name == "Мастера":
        refresh_masters_page()
    elif page_name == "Услуги":
        refresh_services_page()
    elif page_name == "Запись":
        refresh_booking_page()
    elif page_name == "График работы":
        refresh_schedule_page()
    elif page_name == "Главная":
        refresh_main_page()


# --- Стили ---
LARGE_FONT = ('Arial', 20)
MEDIUM_FONT = ('Arial', 16)
SMALL_FONT = ('Arial', 14)
BIG_TIME_FONT = ('Arial', 18, 'bold')

style = ttk.Style()
style.configure("Menu.TButton", font=LARGE_FONT, padding=12)
style.configure("Selected.Menu.TButton", background="#CCCCCC", font=LARGE_FONT, padding=12)
style.configure("Large.TButton", font=MEDIUM_FONT, padding=8)
style.configure("Small.TButton", font=SMALL_FONT)
style.configure("Treeview", font=MEDIUM_FONT)
style.configure("Treeview.Heading", font=MEDIUM_FONT)
style.configure("TLabel", font=MEDIUM_FONT)
style.configure("TEntry", font=MEDIUM_FONT)



def make_large_entry(widget):
    """Настройка шрифта для полей ввода"""
    try:
        widget.configure(font=MEDIUM_FONT)
    except Exception:
        pass


def make_large_combo(combo):
    """Настройка шрифта для комбобоксов"""
    try:
        combo.configure(font=MEDIUM_FONT, width=16)
    except Exception:
        try:
            combo.configure(width=16)
        except Exception:
            pass


# ========== ГЛАВНАЯ СТРАНИЦА ==========
def refresh_main_page():
    """Обновление главной страницы"""
    # Автоматически обновляем статусы записей (если есть такая функция)
    try:
        # Если есть функция для автообновления статусов
        update_appointment_status_based_on_time()
    except:
        pass  # Если функции нет, пропускаем

    # Получаем текущие показатели
    today_str = date.today().strftime('%Y-%m-%d')
    try:
        appointments = get_appointments_by_date(today_str)
    except Exception as e:
        print(f"Ошибка при загрузке статистики: {e}")
        appointments = []

    global stats_data1, stats_data2
    stats_data1[0].configure(text=f"Записей сегодня: {len(appointments)}")
    stats_data1[1].configure(text=f"Выполнено: {len([a for a in appointments if a.get('status') == 'Выполнено'])}")
    stats_data1[2].configure(text=f"Отменено: {len([a for a in appointments if a.get('status') == 'Отменено'])}")

    stats_data2[0].configure(text=f"Ожидается: {len([a for a in appointments if a.get('status') == 'Ожидается'])}")
    stats_data2[1].configure(text=f"В работе: {len([a for a in appointments if a.get('status') == 'В процессе'])}")


    page = pages["Главная"]

    # Очищаем старые данные в таблице
    for item in main_tree.get_children():
        main_tree.delete(item)

    # Получаем сегодняшние записи
    today_str = date.today().strftime('%Y-%m-%d')
    try:
        appointments = get_appointments_by_date(today_str)
    except Exception as e:
        print(f"Ошибка при загрузке записей: {e}")
        appointments = []

    # Добавляем записи в таблицу
    for app in appointments:
        # Проверяем наличие всех необходимых полей
        start_time = timedelta_to_string(app.get('start_time', ''))
        end_time = timedelta_to_string(app.get('end_time', ''))
        master_name = app.get('master_name', 'Не указан')
        service_name = app.get('service_name', 'Не указана')
        client_name = app.get('client_name', 'Не указан')
        status = app.get('status', 'Неизвестно')  # Добавляем статус

        # Определяем цвет строки в зависимости от статуса
        tags = ()
        if status == 'Выполнено':
            tags = ('completed',)
        elif status == 'Отменено':
            tags = ('cancelled',)
        elif status == 'Ожидается':
            tags = ('pending',)
        elif status == 'В процессе':
            tags = ('in_progress',)

        main_tree.insert("", tk.END, values=(
            f"{start_time}-{end_time}",
            master_name,
            service_name,
            client_name,
            status  # Добавляем статус в таблицу
        ), tags=tags)
    # Настраиваем цвета для тегов
    main_tree.tag_configure('completed', background='#e6ffe6')  # Зеленый для выполненных
    main_tree.tag_configure('cancelled', background='#ffe6e6')  # Красный для отмененных
    main_tree.tag_configure('pending', background='#ffffe6')    # Желтый для ожидаемых
    main_tree.tag_configure('in_progress', background='#e6f2ff') # Голубой для в процессе


def create_main_page(page):
    """Создание главной страницы"""
    page.grid_rowconfigure(4, weight=1)
    page.grid_columnconfigure(0, weight=1)

    # Заголовок
    ttk.Label(page, text="Главная", font=LARGE_FONT).grid(row=0, column=0, pady=10, sticky="n")
    ttk.Label(page, text=f"Сегодня: {date.today().strftime('%d.%m.%Y')}", font=MEDIUM_FONT).grid(
        row=1, column=0, sticky="nw", padx=12, pady=5
    )

    # Статистика
    stats_frame = ttk.Frame(page)
    stats_frame.grid(row=2, column=0, sticky="ew", padx=12, pady=(6, 12))
    stats_frame.grid_columnconfigure(0, weight=1)

    ttk.Label(stats_frame, text="Текущие показатели", font=MEDIUM_FONT).grid(row=0, column=0, pady=6, sticky="w")

    row1_bg = "#D3D3D3"
    row2_bg = "#D3D3D3"

    first_row = tk.Frame(stats_frame, bg=row1_bg)
    first_row.grid(row=1, column=0, sticky="ew", padx=20, pady=(2, 0))
    for c in range(3):
        first_row.grid_columnconfigure(c, weight=1, uniform="statscol")

    second_row = tk.Frame(stats_frame, bg=row2_bg)
    second_row.grid(row=2, column=0, sticky="ew", padx=20, pady=(8, 6))
    for c in range(3):
        second_row.grid_columnconfigure(c, weight=1, uniform="statscol")

    global stats_data1, stats_data2
    stats_data1 = [
        tk.Label(first_row, text="", font=SMALL_FONT, bg=row1_bg, anchor="center"),
        tk.Label(first_row, text="", font=SMALL_FONT, bg=row1_bg, anchor="center"),
        tk.Label(first_row, text="", font=SMALL_FONT, bg=row1_bg, anchor="center"),
    ]

    stats_data2 = [
        tk.Label(second_row, text="", font=SMALL_FONT, bg=row2_bg, anchor="center"),
        tk.Label(second_row, text="", font=SMALL_FONT, bg=row2_bg, anchor="center")
    ]

    for col, lbl in enumerate(stats_data1):
        lbl.grid(row=0, column=col, sticky="nsew", padx=6, pady=8)

    for col, lbl in enumerate(stats_data2):
        lbl.grid(row=0, column=col, sticky="nsew", padx=6, pady=8)

    ttk.Label(page, text="Ближайшие записи", font=MEDIUM_FONT).grid(row=3, column=0, pady=10, sticky="w", padx=12)

    # Таблица с записями
    tree_frame = ttk.Frame(page)
    tree_frame.grid(row=4, column=0, sticky="nsew", padx=12, pady=(0, 12))
    tree_frame.grid_rowconfigure(0, weight=1)
    tree_frame.grid_columnconfigure(0, weight=1)

    columns = ("Время", "Мастер", "Услуга", "Клиент")
    global main_tree
    main_tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=8)
    main_tree.grid(row=0, column=0, sticky="nsew")

    # Настройка заголовков
    for col in columns:
        main_tree.heading(col, text=col.capitalize())
        main_tree.column(col, width=150)

    # Scrollbar
    scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=main_tree.yview)
    scrollbar.grid(row=0, column=1, sticky="ns")
    main_tree.configure(yscrollcommand=scrollbar.set)

    # Заполняем данными
    refresh_main_page()


# ========== СТРАНИЦА ЗАПИСИ ==========
def refresh_booking_page():
    """Обновление страницы записи"""
    global selected_master_id_on_booking_page, selected_client_id_on_booking_page, current_masters_on_booking_page, current_clients_on_booking_page

    # Убираем значения со всех полей страницы
    client_combo.set("")
    master_combo.set("")
    service_combo.set("")
    phone_entry.delete(0, tk.END)


    # Получаем всех мастеров и клиентов, если произошла ошибка, замораживаем страницу с бронированием
    try:
        current_masters_on_booking_page = get_id_and_full_name_masters()
        current_clients_on_booking_page = get_clients()
    except Exception:
        disabled_booking_page()
        messagebox.showerror("Ошибка", "Произошла ошибка при подключении к БД.")
        return 0

    # Если списки мастеров или клиентов пусты также замораживаем страницу
    if len(current_masters_on_booking_page) == 0 or len(current_clients_on_booking_page) == 0:
        disabled_booking_page()
        messagebox.showwarning("Внимание", "В системе отсутствуют клиенты (мастера).")
        return 0

    enabled_booking_page()

    # Формируем списки для полей со списком
    master_combobox_values = [f"{master_record['id']}|{master_record["full_name"]}" for master_record in current_masters_on_booking_page]
    client_combobox_values = [f"{client_record['id']}|{client_record['full_name']}" for client_record in current_clients_on_booking_page]

    # Формируем combobox с мастером и клиентом
    master_combo.configure(values=master_combobox_values)
    client_combo.configure(values=client_combobox_values )
    master_combo.set(current_masters_on_booking_page[0]["full_name"])
    client_combo.set(current_clients_on_booking_page[0]["full_name"])

    phone_entry.insert(0, current_clients_on_booking_page[0]["phone"])

    selected_master_id_on_booking_page = int(current_masters_on_booking_page[0]["id"])
    selected_client_id_on_booking_page = int(current_clients_on_booking_page[0]["id"])

    # Загружаем расписание мастера
    load_master_information()


def load_master_information():
    """Обновление расписания мастера"""
    global selected_master_id_on_booking_page, current_services_on_booking_page, selected_service_id_on_booking_page

    # Обновляем расписание выбранного мастера
    update_booking_master_information()

    # Обновляем список услуг для выбранного мастера
    try:
        current_services_on_booking_page = get_services_by_master(selected_master_id_on_booking_page)
    except Exception:
        current_services_on_booking_page = []
        messagebox.showerror("Ошибка", "Произошла ошибка при подключении к БД.")

    # Если мастеру не назначено ни одной услуги
    if len(current_services_on_booking_page) == 0:
        service_combo['values'] = []
        service_combo.set("")
        selected_service_id_on_booking_page = -1
        price_label.config(text="Цена: ____ р")
        messagebox.showwarning("Внимание", "Для данного мастера не определена ни одна услуга.")
        return 0

    service_combo_values = [f"{service_record['id']}|{service_record['name']}|{service_record['duration']} мин" for service_record in current_services_on_booking_page]
    service_combo['values'] = service_combo_values

    service_combo.set(current_services_on_booking_page[0]["name"])
    selected_service_id_on_booking_page = int(current_services_on_booking_page[0]["id"])
    price_label.config(text=f"Цена: {current_services_on_booking_page[0]["cost"]} р")


def update_booking_master_information():
    """Обновление ограничений по дате"""
    global selected_master_id_on_booking_page, booking_master_schedule_table
    booking_master_schedule_table.deselect(row="all")
    current_day_booking_table.deselect(row="all")

    # Получаем выбранную дату из календаря
    selected_date = booking_page_date_calendar.get_date()

    # Получаем расписание мастера на эту дату
    try:
        schedule = get_master_schedule_for_date(selected_master_id_on_booking_page, selected_date)
        day_appointments = get_appointments_by_date_and_master(selected_date, selected_master_id_on_booking_page)
    except Exception:
        messagebox.showerror("Ошибка", "Произошла ошибка при подключении к БД.")
        return 0


    if schedule:
        # Преобразуем время в строку
        start_time = schedule['start_time']
        end_time = schedule['end_time']

        start_time_string = timedelta_to_string(start_time)
        end_time_string = timedelta_to_string(end_time)

        date_label.config(text=f"Время работы: {start_time_string} - {end_time_string}")
    else:
        date_label.config(text="Мастер не работает в этот день")

    # Получаемем выбранный день из колендаря
    selected_date = booking_page_date_calendar.get_date()
    # Получаем доступные слоты
    time_slots = get_available_time_slots(selected_master_id_on_booking_page, selected_date)

    table_slot_data = []

    # Заполняем данные таблицы
    for slot in time_slots:
        table_slot_data.append([slot['start'], slot['end'], f"Свободное время: {slot['start']} - {slot['end']}"])
    booking_master_schedule_table.set_sheet_data(table_slot_data)


    appointment_table_data = []

    for appointment in day_appointments:
        appointment_info = f"Запись №{appointment["id"]} на {timedelta_to_string(appointment["start_time"])}-{timedelta_to_string(appointment["end_time"])}\nКлиент - {appointment["client_name"]}: {appointment["service_name"]}\nСтатус: {appointment["status"]}"
        appointment_table_data.append([appointment["id"], appointment["status"], appointment_info])

    current_day_booking_table.set_sheet_data(appointment_table_data)




def confirm_booking():
    """Подтверждение записи"""

    # Собираем id клиента мастера и записи, загрудаем дату из календаря
    global selected_master_id_on_booking_page, selected_client_id_on_booking_page, selected_service_id_on_booking_page, current_services_on_booking_page
    selected_date = booking_page_date_calendar.get_date()

    # Проверяем id сервиса, только он может быть неправильный
    if selected_service_id_on_booking_page == -1:
        messagebox.showerror("Ошибка", "Услуга не выбрана!")
        return 0


    # Получаем время услуги
    service_time = int(list(filter(lambda service_record: service_record["id"] == selected_service_id_on_booking_page, current_services_on_booking_page))[0]["duration"])

    # Выбираем все выбранные временные промежутки
    selected_time_cells = booking_master_schedule_table.get_selected_cells()
    selected_time_zone = []
    for selected_cells in selected_time_cells:
        #Выбираем из таблицы начало и конец временной ячейки
        cells_row = selected_cells[0]
        start_time, end_time, _ = booking_master_schedule_table.get_row_data(r=cells_row)
        selected_time_zone.append([start_time, end_time])

    # Сортируем временные ячейки по дате начала
    selected_time_zone.sort(key=lambda zone: zone[0])
    booking_start_time_str = ""
    booking_end_time_str = ""

    # Проходимся по всем ячейкам
    for i, time_zone in enumerate(selected_time_zone):
        start_time, end_time = time_zone
        # Если ячейка первая - ее начало - начало записи
        if i == 0:
            booking_start_time_str = start_time
        # Если ячейка последняя - ее конец - конец записи
        if i == len(selected_time_zone) - 1:
            booking_end_time_str = end_time
        # Если конец текущей ячейки не равен началу последней - они выбраны не по порядку - вызываем ошибку
        if (i != len(selected_time_zone) - 1) and (end_time != selected_time_zone[i+1][0]):
            messagebox.showerror("Ошибка", "При выборе нескольких временных промежутков они должны следовать друг за другом.")
            return 0

    # Если кол-во минут на услугу больше, чем выбранное, вызываем ошибку
    if service_time > len(selected_time_zone) * 30:
        messagebox.showerror("Ошибка", "Время выполнения услуги превышает выбранный промежуток.")
        return 0

    # Переводим в необходимый формат
    booking_start_time = datetime.strptime(booking_start_time_str, '%H:%M')
    booking_end_time = datetime.strptime(booking_end_time_str, '%H:%M')

    # Добавляем запись в БД
    try:
        add_appointment(
            client_id=selected_client_id_on_booking_page,
            master_id=selected_master_id_on_booking_page,
            service_id=selected_service_id_on_booking_page,
            appointment_date=selected_date,
            start_time=booking_start_time,
            end_time=booking_end_time
        )
        messagebox.showinfo("Успех", "Запись успешно создана!")

        # Сбрасываем форму
        refresh_booking_page()  # Полностью обновляем страницу

    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось создать запись: {str(e)}")


def cancel_booking():
    """Отмена выбранной записи"""
    selected_appointment_cells = tuple(current_day_booking_table.get_selected_cells())

    if len(selected_appointment_cells) == 0:
        messagebox.showerror("Ошибка", "Выберите запись для отмены на главной странице!")
        return 0

    selected_row = selected_appointment_cells[0][0]
    appointment_id, appointment_status, _ = current_day_booking_table.get_row_data(r=selected_row)

    if appointment_status != 'Ожидается':
        messagebox.showerror("Ошибка", "Нельзя отменить уже начатые, завершенные или отмененные записи!")
        return 0

    if not messagebox.askyesno("Подтверждение", "Вы уверены, что хотите отменить эту запись?"):
        return 0

    try:
        # Изменяем статус записи на "Отменено"
        cancel_appointment(appointment_id)
        messagebox.showinfo("Успех", "Запись успешно отменена!")
        refresh_booking_page()  # Обновляем главную страницу
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при отмене записи: {str(e)}")


def create_booking_page(page): #-------------------------------------------------------------------------------Тут воняет сильнее---------------------------------------------------------
    """Создание страницы записи"""
    global selected_time_slot, selected_master_id_on_booking_page
    selected_time_slot = None
    selected_master_id_on_booking_page = None

    page.grid_rowconfigure(4, weight=0)
    page.grid_rowconfigure(3, weight=0)
    page.grid_rowconfigure(2, weight=1)
    page.grid_columnconfigure(0, weight=1)

    # Заголовок
    ttk.Label(page, text="Запись", font=LARGE_FONT).grid(row=0, column=0, pady=10, sticky="n")

    # Верхняя панель с выбором даты и мастера
    top_frame = ttk.Frame(page)
    top_frame.grid(row=1, column=0, sticky="ew", pady=(0, 10), padx=12)
    top_frame.grid_columnconfigure(0, weight=1)

    # Дата
    date_frame = ttk.Frame(top_frame)
    date_frame.grid(row=0, column=0, sticky="w", pady=4)
    ttk.Label(date_frame, text="Дата:", font=MEDIUM_FONT).pack(side="left")

    global booking_page_date_calendar

    booking_page_date_calendar = DateEntry(date_frame, date_pattern='dd.mm.yyyy', state='readonly')
    booking_page_date_calendar.pack(side="left", padx=6)
    make_large_entry(booking_page_date_calendar)

    # Обработчик изменения даты
    def on_date_selected(event):
        update_booking_master_information()

    booking_page_date_calendar.bind("<<DateEntrySelected>>", on_date_selected)

    # Мастер
    master_frame = ttk.Frame(top_frame)
    master_frame.grid(row=1, column=0, sticky="w", pady=4)
    ttk.Label(master_frame, text="Мастер:", font=MEDIUM_FONT).pack(side="left")

    global master_combo
    master_combo = ttk.Combobox(master_frame, state="readonly")
    master_combo.pack(side="left", padx=6)
    master_combo.bind("<<ComboboxSelected>>", lambda event: master_combobox_selected_function(event))
    make_large_combo(master_combo)

    # Метка с информацией о времени работы
    global date_label
    date_label = ttk.Label(top_frame, text="Выберите мастера и дату", font=SMALL_FONT, foreground="blue")
    date_label.grid(row=2, column=0, sticky="w", pady=4)

    # Центральная область с доступным временем
    center_container = ttk.Frame(page)
    center_container.grid(row=2, column=0, sticky="nsew", pady=6, padx=12)
    center_container.grid_rowconfigure(0, weight=1)
    center_container.grid_columnconfigure(0, weight=1)

    tables_notebook = ttk.Notebook(master=center_container)
    tables_notebook.grid(row=0, column=0, sticky="nsew")

    schedule_table_frame = tk.Frame(tables_notebook)
    schedule_table_frame.grid_columnconfigure(0, weight=1)
    schedule_table_frame.grid_rowconfigure(0, weight=1)
    tables_notebook.add(schedule_table_frame, text="Свободное время на текущий день")

    appointment_table_frame = tk.Frame(tables_notebook)
    appointment_table_frame.grid_columnconfigure(0, weight=1)
    appointment_table_frame.grid_rowconfigure(0, weight=1)
    tables_notebook.add(appointment_table_frame, text="Записи на текущий день")

    global booking_master_schedule_table, current_day_booking_table, root

    booking_master_schedule_table = Sheet(
        schedule_table_frame,
        show_x_scrollbar=False,
        empty_vertical=False,
        empty_horizontal=False,
        show_top_left=False,
        show_row_index=False,
        headers=["", "", "Свободное время на текущий день"],
        total_columns=3,

        default_row_height="4",
        header_font=("Arial", 20, "normal"),
        font=("Arial", 18, "normal"),

        header_selected_cells_bg="#ffffff",
    )
    booking_master_schedule_table.hide_columns(columns=[0, 1])



    booking_master_schedule_table.grid(row=0, column=0, sticky="nsew")
    booking_master_schedule_table.set_options(auto_resize_columns=1)
    booking_master_schedule_table.enable_bindings("toggle_select")

    current_day_booking_table = Sheet(
        appointment_table_frame,
        show_x_scrollbar=False,
        empty_vertical=False,
        empty_horizontal=False,
        show_top_left=False,
        show_row_index=False,
        headers=["", "", "Записи на текущий день"],
        total_columns=2,

        default_row_height="4",
        header_font=("Arial", 20, "normal"),
        font=("Arial", 14, "normal"),

        header_selected_cells_bg="#ffffff",
    )
    current_day_booking_table.hide_columns(columns=[0, 1])
    current_day_booking_table.grid(row=0, column=0, sticky="nsew")
    current_day_booking_table.set_options(auto_resize_columns=1)
    current_day_booking_table.enable_bindings("single_select")



    # Информация о клиенте и услуге
    info_frame = ttk.Frame(page, relief="groove", borderwidth=1, padding=10)
    info_frame.grid(row=3, column=0, sticky="ew", pady=(10, 0), padx=12)
    info_frame.grid_columnconfigure(0, weight=0)
    info_frame.grid_columnconfigure(1, weight=1)
    info_frame.grid_columnconfigure(2, weight=0)
    info_frame.grid_columnconfigure(3, weight=1)
    info_frame.grid_columnconfigure(4, weight=0)

    # Клиент
    ttk.Label(info_frame, text="Клиент:", font=MEDIUM_FONT).grid(row=0, column=0, sticky="w", pady=4, padx=(0, 4))

    global client_combo
    client_combo = ttk.Combobox(info_frame, state="readonly")
    client_combo.grid(row=0, column=1, sticky="ew", pady=4, padx=(0, 12))
    client_combo.bind("<<ComboboxSelected>>", lambda event: client_combobox_selected_function(event))
    make_large_combo(client_combo)

    # Телефон
    ttk.Label(info_frame, text="Телефон:", font=MEDIUM_FONT).grid(row=0, column=2, sticky="w", pady=4, padx=(0, 4))

    global phone_entry
    phone_entry = ttk.Entry(info_frame)
    phone_entry.grid(row=0, column=3, sticky="ew", pady=4, padx=(0, 12))
    make_large_entry(phone_entry)

    # Кнопка "Новый клиент"
    ttk.Button(info_frame, text="Новый клиент", style="Large.TButton",
               command=lambda: show_page("Клиенты")).grid(row=0, column=4, pady=4)

    # Услуга
    ttk.Label(info_frame, text="Услуга:", font=MEDIUM_FONT).grid(row=1, column=0, sticky="w", pady=8, padx=(0, 4))

    global service_combo
    service_combo = ttk.Combobox(info_frame, state="readonly")
    service_combo.grid(row=1, column=1, sticky="ew", pady=8, padx=(0, 12))
    service_combo.bind("<<ComboboxSelected>>", lambda event: service_combobox_selected_function(event))
    make_large_combo(service_combo)

    # Цена
    global price_label
    price_label = ttk.Label(info_frame, text="Цена: ____ р", font=MEDIUM_FONT)
    price_label.grid(row=1, column=2, columnspan=2, sticky="w", pady=8)

    # Кнопки действий
    action_frame = ttk.Frame(page)
    action_frame.grid(row=4, column=0, pady=(20, 10), padx=12, sticky="ew")
    action_frame.grid_columnconfigure(0, weight=1)
    action_frame.grid_columnconfigure(1, weight=1)

    # Все кнопки
    global access_booking_button, delete_booking_button
    access_booking_button = ttk.Button(action_frame, text="Подтвердить запись", style="Large.TButton", command=confirm_booking)
    access_booking_button.grid(row=0, column=0, padx=8, sticky="ew")

    delete_booking_button = ttk.Button(action_frame, text="Удалить", style="Large.TButton", command=cancel_booking)
    delete_booking_button.grid(row=0, column=1, padx=8, sticky="ew")


def master_combobox_selected_function(event):
    """Выбор мастера через combobox"""
    global selected_master_id_on_booking_page
    selected_master_id_on_booking_page, master_full_name = master_combo.get().split("|")
    selected_master_id_on_booking_page = int(selected_master_id_on_booking_page)
    master_combo.set(master_full_name)
    load_master_information()


def client_combobox_selected_function(event):
    """Выбор клиента через combobox"""
    global selected_client_id_on_booking_page, current_clients_on_booking_page
    selected_client_id_on_booking_page, client_full_name = client_combo.get().split("|")
    selected_client_id_on_booking_page = int(selected_client_id_on_booking_page)
    client_combo.set(client_full_name)


    selected_client_phone = list(filter(lambda client_record: client_record["id"] == selected_client_id_on_booking_page, current_clients_on_booking_page))[0]["phone"]
    phone_entry.delete(0, tk.END)
    phone_entry.insert(0, selected_client_phone)


def service_combobox_selected_function(event):
    """Выбор сервиса"""
    global current_services_on_booking_page, selected_service_id_on_booking_page

    selected_service_id_on_booking_page, service_name, _ = service_combo.get().split("|")
    selected_service_id_on_booking_page = int(selected_service_id_on_booking_page)
    service_combo.set(service_name)

    service_price = list(filter(lambda service_record: service_record["id"] == selected_service_id_on_booking_page, current_services_on_booking_page))[0]["cost"]
    price_label.config(text=f"Цена: {service_price} р")


def disabled_booking_page():
    """Функция, используемая для заморозки страницы бронирования, если нету ни одного мастера или клиента"""
    booking_page_date_calendar.configure(state="disabled")
    master_combo.configure(state="disabled")
    client_combo.configure(state="disabled")
    phone_entry.configure(state="disabled")
    service_combo.configure(state="disabled")

    access_booking_button.configure(state="disabled")
    delete_booking_button.configure(state="disabled")


def enabled_booking_page():
    """Функция, используемая для разморозки страницы бронирования"""
    booking_page_date_calendar.configure(state="readonly")
    master_combo.configure(state="readonly")
    client_combo.configure(state="readonly")
    phone_entry.configure(state="enabled")
    service_combo.configure(state="readonly")

    access_booking_button.configure(state="enabled")
    delete_booking_button.configure(state="enabled")


# ========== СТРАНИЦА ГРАФИКА РАБОТЫ ==========
def refresh_schedule_page():
    """Обновление страницы графика работы"""
    # Обновляем список мастеров
    masters = get_masters()
    master_names = [f"{m['surname']} {m['name']}" for m in masters]
    master_combo_sched['values'] = master_names
    if master_names:
        master_combo_sched.set(master_names[0])
        load_master_schedule()


def load_master_schedule():
    """Загрузка расписания мастера"""
    try:
        # Получаем ID выбранного мастера

        selected_master = master_combo_sched.get()
        if not selected_master:
            return 0

        master_surname, master_name = selected_master.split()
        master_id = get_master_id_by_name(master_surname, master_name)
        if not master_id:
            return 0

        # Получаем выбранную дату из календаря
        selected_schedule_day = shedule_week_calendar.get()

        # Получаем начало недели
        week_start = datetime.strptime(selected_schedule_day, '%d.%m.%Y').date()
        week_start -= timedelta(days=week_start.weekday())

        # Вычисляем конец недели (воскресенье)
        week_end = week_start + timedelta(days=6)

        # Обновляем метку с информацией о неделе
        week_info_label.config(text=f"Неделя {week_start.strftime('%d.%m.%Y')} - {week_end.strftime('%d.%m.%Y')}")

        # Загружаем расписание на всю неделю
        try:
            schedule = get_master_schedule(master_id, week_start)
        except Exception as e:
            print(f"Ошибка загрузки расписания: {e}")
            schedule = []

        # Обновляем данные в таблице для всех 7 дней недели
        for i in range(7):
            current_date = week_start + timedelta(days=i)

            # Обновляем метку с датой
            day_labels[i].config(text=f"{current_date.strftime('%d.%m')} ({['Пн', 'Вт', 'Ср', 'Чт', 'Пт', 'Сб', 'Вс'][i]})")

            # Ищем запись на эту дату
            current_date_schedule = tuple(filter(lambda schedule_record: schedule_record['work_date'] == current_date, schedule))
            day_schedule = current_date_schedule[0] if len(current_date_schedule) == 1 else None

            # Устанавливаем значения

            if day_schedule is not None:
                day_vars[i].set(1)
                # Извлекаем время начала и окончания
                start_time = day_schedule['start_time']
                end_time = day_schedule['end_time']

                # ПРЕОБРАЗУЕМ В СТРОКУ ПРАВИЛЬНО
                start_time_str = timedelta_to_string(start_time)
                end_time_str = timedelta_to_string(end_time)

                start_combos[i].set(start_time_str)
                end_combos[i].set(end_time_str)
            else:
                # По умолчанию: будни - работает, выходные - нет
                day_vars[i].set(0)
                start_combos[i].set("Нет данных")
                end_combos[i].set("Нет данных")

    except Exception as e:
        print(f"Общая ошибка в load_master_schedule: {e}")
        messagebox.showerror("Ошибка", f"Не удалось загрузить расписание: {str(e)}")


def save_schedule():
    """Сохранение расписания мастера"""
    try:
        # Получаем ID мастера
        selected_master = master_combo_sched.get()
        if not selected_master:
            messagebox.showerror("Ошибка", "Выберите мастера!")
            return 0

        master_surname, master_name = selected_master.split()
        master_id = get_master_id_by_name(master_surname, master_name)
        if not master_id:
            messagebox.showerror("Ошибка", "Мастер не найден!")
            return 0


        # Получаем выбранную неделю из календаря

        # Получаем дату из DateEntry
        selected_date_str = shedule_week_calendar.get()

        # Преобразуем строку в дату
        week_start = datetime.strptime(selected_date_str, '%d.%m.%Y').date()

        # Находим понедельник
        week_start -= timedelta(days=week_start.weekday())
        week_end = week_start + timedelta(days=6)

        # Массив существующих записей на текущую неделю
        week_appointments = get_week_appointments(week_start, week_end)

        # Собираем данные для всех дней недели
        day_data = []

        for i in range(7):
            current_date = week_start + timedelta(days=i)

            current_date_appointments = list(filter(
                lambda current_appointment: current_appointment["appointment_date"] == current_date,
                week_appointments)
            )


            if day_vars[i].get() == 1:  # Если мастер работает в этот день
                # Проверяем время
                start_time_string = start_combos[i].get()
                end_time_string = end_combos[i].get()

                # Проверяем заполнение полей рабочих дней сотрудника
                if start_time_string == "Нет данных" or end_time_string == "Нет данных":
                    messagebox.showerror(
                        "Ошибки в данных",
                        f"Укажите время работы для {current_date.strftime('%d.%m.%Y')}"
                    )
                    return 0

                start_time = datetime.strptime(start_time_string, '%H:%M')
                end_time = datetime.strptime(end_time_string, '%H:%M')

                # Проверяем, что начало раньше конца
                if start_time >= end_time:
                    messagebox.showerror(
                        "Ошибки в данных",
                        f"Время начала должно быть раньше времени окончания для {current_date.strftime('%d.%m.%Y')}"
                    )
                    return 0

                for current_appointment in current_date_appointments:
                    start_appointment_time = datetime.strptime(timedelta_to_string(current_appointment["start_time"]), '%H:%M')
                    end_appointment_time = datetime.strptime(timedelta_to_string(current_appointment["end_time"]),'%H:%M')

                    if start_time > start_appointment_time or end_time < end_appointment_time:
                        messagebox.showerror(
                            "Ошибка в данных",
                            f"Для дня {current_date} существуют записи, выходящие за пределы указанного времени!"
                        )
                        return 0



                day_data.append({
                    'day_of_week': i + 1,  # 1=Пн, 7=Вс
                    'work_date': current_date,
                    'start_time': start_time,
                    'end_time': end_time
                })

        # Сохраняем в БД
        try:
            set_master_schedule(master_id, week_start, week_end, day_data)
            messagebox.showinfo("Успех", f"Расписание на неделю сохранено!")

            # Перезагружаем расписание для отображения изменений
            load_master_schedule()

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить расписание в базу данных: {str(e)}")

    except Exception as e:
        messagebox.showerror("Ошибка", f"Неожиданная ошибка: {str(e)}")



def create_schedule_page(page):
    """Создание страницы графика работы"""
    page.grid_rowconfigure(1, weight=1)
    page.grid_columnconfigure(0, weight=1)

    # Заголовок
    ttk.Label(page, text="График работы", font=LARGE_FONT).grid(row=0, column=0, pady=12, sticky="n")

    main_frame = ttk.Frame(page)
    main_frame.grid(row=1, column=0, sticky="nsew", padx=12, pady=6)
    main_frame.grid_columnconfigure(0, weight=1)

    # Верхняя панель
    top_frame = ttk.Frame(main_frame)
    top_frame.grid(row=0, column=0, sticky="ew", pady=6)

    ttk.Label(top_frame, text="Мастер:", font=MEDIUM_FONT).grid(row=0, column=0, sticky="w")

    global master_combo_sched
    master_combo_sched = ttk.Combobox(top_frame, state="readonly")
    master_combo_sched.grid(row=0, column=1, sticky="w", padx=8)
    master_combo_sched.bind("<<ComboboxSelected>>", lambda e: load_master_schedule())
    make_large_combo(master_combo_sched)

    ttk.Label(top_frame, text="Неделя с:", font=MEDIUM_FONT).grid(row=0, column=2, sticky="w", padx=(12, 0))

    global shedule_week_calendar
    shedule_week_calendar = DateEntry(top_frame, width=14, date_pattern='dd.mm.yyyy', state='readonly')
    shedule_week_calendar.grid(row=0, column=3, sticky="w", padx=8)
    shedule_week_calendar.bind("<<DateEntrySelected>>", lambda e: load_master_schedule())
    make_large_entry(shedule_week_calendar)

    # Кнопки навигации по неделям
    nav_frame = ttk.Frame(top_frame)
    nav_frame.grid(row=0, column=4, sticky="w", padx=(12, 0))


    # Метка с информацией о неделе
    global week_info_label
    week_info_label = ttk.Label(top_frame, text="", font=SMALL_FONT, foreground="blue")
    week_info_label.grid(row=1, column=0, columnspan=5, sticky="w", pady=4)

    # Таблица расписания
    schedule_frame = ttk.Frame(main_frame, relief="groove", borderwidth=1, padding=10)
    schedule_frame.grid(row=1, column=0, sticky="ew", pady=6)

    for c in range(1, 8):
        schedule_frame.grid_columnconfigure(c, weight=1, uniform="schedcol")

    # Создаем глобальные переменные для хранения данных
    global day_vars, start_combos, end_combos, day_labels
    day_vars = []
    start_combos = []
    end_combos = []
    day_labels = []

    time_options = [f"{h:02d}:00" for h in range(8, 23)]

    # Заголовки столбцов
    days_of_week = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Вс"]
    ttk.Label(schedule_frame, text="", font=SMALL_FONT).grid(row=0, column=0)
    for i, day in enumerate(days_of_week):
        day_label = ttk.Label(schedule_frame, text=day, font=MEDIUM_FONT)
        day_label.grid(row=0, column=i + 1, sticky="nsew", padx=4, pady=6)
        day_labels.append(day_label)

    # Работает
    ttk.Label(schedule_frame, text="Работает", font=MEDIUM_FONT).grid(row=1, column=0, sticky="w", padx=6)
    for i in range(7):
        var = tk.IntVar(value=1 if i < 5 else 0)
        day_vars.append(var)
        chk = tk.Checkbutton(schedule_frame, variable=var, onvalue=1, offvalue=0,
                             font=MEDIUM_FONT, text="", padx=6, pady=4)
        chk.grid(row=1, column=i + 1, sticky="nsew", padx=4, pady=6)

    # Начало работы
    ttk.Label(schedule_frame, text="Начало", font=MEDIUM_FONT).grid(row=2, column=0, sticky="w", padx=6, pady=6)
    for i in range(7):
        combo = ttk.Combobox(schedule_frame, values=time_options, state="readonly", width=10)
        combo.set("09:00")
        combo.grid(row=2, column=i + 1, sticky="nsew", padx=4, pady=6)
        start_combos.append(combo)
        make_large_combo(combo)

    # Конец работы
    ttk.Label(schedule_frame, text="Конец", font=MEDIUM_FONT).grid(row=3, column=0, sticky="w", padx=6, pady=6)
    for i in range(7):
        combo = ttk.Combobox(schedule_frame, values=time_options, state="readonly", width=10)
        combo.set("18:00")
        combo.grid(row=3, column=i + 1, sticky="nsew", padx=4, pady=6)
        end_combos.append(combo)
        make_large_combo(combo)

    # Кнопка сохранения
    btn_frame = ttk.Frame(main_frame)
    btn_frame.grid(row=2, column=0, sticky="w", pady=6)
    ttk.Button(btn_frame, text="Сохранить расписание", style="Large.TButton",
               command=save_schedule).pack(side="left", padx=4)

    # Загружаем данные при открытии страницы
    refresh_schedule_page()



# ========== СТРАНИЦА УСЛУГ ==========
def refresh_services_page():
    """Обновление страницы услуг"""
    # Очищаем таблицу
    for item in services_tree.get_children():
        services_tree.delete(item)

    # Загружаем услуги из БД
    services = get_services()
    for service in services:
        services_tree.insert("", tk.END, values=(
            service['id'],  # Скрытый ID
            service['name'],
            f"{service['duration']} мин",
            f"{service['cost']} р",
        ))

    # Скрываем первый столбец с ID
    services_tree.column("#1", width=0, stretch=False)
    services_tree.heading("#1", text="")


def on_service_select(event):
    """Обработка выбора услуги в таблице"""
    selection = services_tree.selection()
    if selection:
        item = services_tree.item(selection[0])
        values = item['values']

        # Заполняем поля формы данными выбранной услуги
        if len(values) >= 4:
            # ID в values[0] (скрытый)
            service_name_entry.delete(0, tk.END)
            service_name_entry.insert(0, values[1])

            # Извлекаем число из строки "60 мин"
            duration_str = values[2].replace(' мин', '').strip()
            service_duration_entry.delete(0, tk.END)
            service_duration_entry.insert(0, duration_str)

            # Извлекаем число из строки "1000 р"
            cost_str = values[3].replace(' р', '').strip()
            service_cost_entry.delete(0, tk.END)
            service_cost_entry.insert(0, cost_str)



def edit_service():
    """Редактирование выбранной услуги"""
    selection = services_tree.selection()
    if not selection:
        messagebox.showerror("Ошибка", "Выберите услугу для редактирования!")
        return

    item = services_tree.item(selection[0])
    values = item['values']

    if not values or len(values) < 4:
        messagebox.showerror("Ошибка", "Не удалось получить данные услуги!")
        return

    service_id = values[0]  # Скрытый ID

    # Получаем данные из формы
    name = service_name_entry.get()
    duration = service_duration_entry.get()
    cost = service_cost_entry.get()


    if not all([name, duration, cost]):
        messagebox.showerror("Ошибка", "Заполните все поля!")
        return

    # Проверяем числовые значения
    try:
        duration_int = int(duration)
        cost_int = int(cost)
    except ValueError:
        messagebox.showerror("Ошибка", "Длительность и стоимость должны быть числами!")
        return

    # Обновляем услугу
    try:
        update_service(service_id, name, duration_int, cost_int)

        messagebox.showinfo("Успех", "Услуга обновлена!")
        refresh_services_page()

        # Очищаем форму
        service_name_entry.delete(0, tk.END)
        service_duration_entry.delete(0, tk.END)
        service_cost_entry.delete(0, tk.END)

    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось обновить услугу: {str(e)}")


def delete_service_from_db():
    """Удаление выбранной услуги"""
    selection = services_tree.selection()
    if not selection:
        messagebox.showerror("Ошибка", "Выберите услугу для удаления!")
        return

    item = services_tree.item(selection[0])
    values = item['values']

    if not values or len(values) < 4:
        messagebox.showerror("Ошибка", "Не удалось получить данные услуги!")
        return

    service_id = values[0]  # Скрытый ID

    # Подтверждение удаления
    if not messagebox.askyesno("Подтверждение", "Вы уверены, что хотите удалить эту услугу?"):
        return

    # Удаляем услугу
    try:
        result = delete_service(service_id)
        if result > 0:
            messagebox.showinfo("Успех", "Услуга удалена!")
            refresh_services_page()

            # Очищаем форму
            service_name_entry.delete(0, tk.END)
            service_duration_entry.delete(0, tk.END)
            service_cost_entry.delete(0, tk.END)

        else:
            messagebox.showerror("Ошибка", "Не удалось удалить услугу!")

    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось удалить услугу: {str(e)}")



def add_service_from_form():
    """Добавление услуги из формы"""
    name = service_name_entry.get()
    duration = service_duration_entry.get()
    cost = service_cost_entry.get()


    if not all([name, duration, cost]):
        messagebox.showerror("Ошибка", "Заполните все поля!")
        return


    # Проверяем числовые значения
    try:
        duration_int = int(duration)
        cost_int = int(cost)
    except ValueError:
        messagebox.showerror("Ошибка", "Длительность и стоимость должны быть числами!")
        return

    # Добавляем услугу
    try:
        add_service(name, duration_int, cost_int)
        messagebox.showinfo("Успех", "Услуга добавлена!")

        # Очищаем форму
        service_name_entry.delete(0, tk.END)
        service_duration_entry.delete(0, tk.END)
        service_cost_entry.delete(0, tk.END)

        # Обновляем таблицу
        refresh_services_page()

    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось добавить услугу: {str(e)}")


def create_services_page(page):
    """Создание страницы услуг"""
    page.grid_rowconfigure(2, weight=1)
    page.grid_columnconfigure(0, weight=1)

    # Заголовок
    ttk.Label(page, text="Услуги", font=LARGE_FONT).grid(row=0, column=0, pady=10, sticky="n")

    # Форма добавления/редактирования
    add_frame = ttk.Frame(page, padding=8, relief="groove", borderwidth=1)
    add_frame.grid(row=1, column=0, sticky="nsew", padx=12, pady=6)
    add_frame.grid_columnconfigure(0, weight=0)
    add_frame.grid_columnconfigure(1, weight=1)

    ttk.Label(add_frame, text="Название услуги:", font=MEDIUM_FONT).grid(row=0, column=0, sticky="w", padx=4, pady=4)
    global service_name_entry
    service_name_entry = ttk.Entry(add_frame)
    service_name_entry.grid(row=0, column=1, sticky="ew", padx=4, pady=4)
    make_large_entry(service_name_entry)

    ttk.Label(add_frame, text="Длительность (мин):", font=MEDIUM_FONT).grid(row=1, column=0, sticky="w", padx=4, pady=4)
    global service_duration_entry
    service_duration_entry = ttk.Entry(add_frame)
    service_duration_entry.grid(row=1, column=1, sticky="ew", padx=4, pady=4)
    make_large_entry(service_duration_entry)

    ttk.Label(add_frame, text="Стоимость:", font=MEDIUM_FONT).grid(row=2, column=0, sticky="w", padx=4, pady=4)
    global service_cost_entry
    service_cost_entry = ttk.Entry(add_frame)
    service_cost_entry.grid(row=2, column=1, sticky="ew", padx=4, pady=4)
    make_large_entry(service_cost_entry)



    # Кнопки формы
    form_buttons_frame = ttk.Frame(add_frame)
    form_buttons_frame.grid(row=4, column=0, columnspan=2, pady=6)

    ttk.Button(form_buttons_frame, text="Добавить", style="Large.TButton",
               command=add_service_from_form).pack(side="left", padx=4)
    ttk.Button(form_buttons_frame, text="Очистить", style="Large.TButton",
               command=lambda: [service_name_entry.delete(0, tk.END),
                                service_duration_entry.delete(0, tk.END),
                                service_cost_entry.delete(0, tk.END),
                                ]).pack(side="left", padx=4)

    # Таблица услуг
    tree_frame = ttk.Frame(page)
    tree_frame.grid(row=2, column=0, sticky="nsew", padx=12, pady=(6, 0))
    tree_frame.grid_rowconfigure(0, weight=1)
    tree_frame.grid_columnconfigure(0, weight=1)

    # Добавляем столбец для ID (будет скрыт)
    columns = ("id", "name", "duration", "cost")
    global services_tree
    services_tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=15)
    services_tree.grid(row=0, column=0, sticky="nsew")

    # Настройка столбцов (первый скрыт)
    services_tree.heading("name", text="Название")
    services_tree.heading("duration", text="Длительность")
    services_tree.heading("cost", text="Стоимость")


    services_tree.column("name", width=200)
    services_tree.column("duration", width=120)
    services_tree.column("cost", width=120)


    # Привязываем обработчик выбора
    services_tree.bind("<<TreeviewSelect>>", on_service_select)

    # Scrollbar
    scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=services_tree.yview)
    scrollbar.grid(row=0, column=1, sticky="ns")
    services_tree.configure(yscrollcommand=scrollbar.set)

    # Кнопки действий под таблицей
    action_frame = ttk.Frame(page)
    action_frame.grid(row=3, column=0, pady=10, padx=12, sticky="ew")
    action_frame.grid_columnconfigure(0, weight=1)
    action_frame.grid_columnconfigure(1, weight=1)

    ttk.Button(action_frame, text="Изменить", style="Large.TButton",
               command=edit_service).grid(row=0, column=0, padx=4, sticky="ew")
    ttk.Button(action_frame, text="Удалить", style="Large.TButton",
               command=delete_service_from_db).grid(row=0, column=1, padx=4, sticky="ew")

    # Загружаем данные
    refresh_services_page()


# ========== СТРАНИЦА КЛИЕНТОВ ==========
def refresh_clients_page():
    """Обновление страницы клиентов"""
    # Очищаем таблицу
    for item in clients_tree.get_children():
        clients_tree.delete(item)

    # Загружаем клиентов из БД
    clients = get_clients()
    for client in clients:
        clients_tree.insert("", tk.END, values=(
            client['id'],  # Скрытый ID
            client['full_name'],
            client['phone']
        ))

    # Скрываем первый столбец с ID
    clients_tree.column("#1", width=0, stretch=False)
    clients_tree.heading("#1", text="")


def on_client_select(event):
    """Обработка выбора клиента в таблице"""
    selection = clients_tree.selection()
    if selection:
        item = clients_tree.item(selection[0])
        values = item['values']

        # Заполняем поля формы данными выбранного клиента
        if len(values) >= 3:
            client_name_entry.delete(0, tk.END)
            client_name_entry.insert(0, values[1])

            client_phone_entry.delete(0, tk.END)
            client_phone_entry.insert(0, values[2])


def edit_client():
    """Редактирование выбранного клиента"""
    selection = clients_tree.selection()
    if not selection:
        messagebox.showerror("Ошибка", "Выберите клиента для редактирования!")
        return

    item = clients_tree.item(selection[0])
    values = item['values']

    if not values or len(values) < 3:
        messagebox.showerror("Ошибка", "Не удалось получить данные клиента!")
        return

    client_id = values[0]  # Скрытый ID

    # Получаем данные из формы
    name = client_name_entry.get()
    phone = client_phone_entry.get()

    if not name or not phone:
        messagebox.showerror("Ошибка", "Заполните все поля!")
        return

    # Проверяем формат телефона
    if not validate_phone(phone):
        messagebox.showerror("Ошибка", "Неверный формат телефона! Должен быть: + и 11 цифр")
        return

    # Разбиваем имя на части
    name_parts = name.split()
    surname = name_parts[0] if len(name_parts) > 0 else ""
    first_name = name_parts[1] if len(name_parts) > 1 else name_parts[0] if name_parts else ""
    patronymic = name_parts[2] if len(name_parts) > 2 else ""

    # Обновляем клиента
    try:
        result = update_client(client_id, surname, first_name, patronymic, phone)
        if result > 0:
            messagebox.showinfo("Успех", "Клиент обновлен!")
            refresh_clients_page()

            # Очищаем форму
            client_name_entry.delete(0, tk.END)
            client_phone_entry.delete(0, tk.END)
        else:
            messagebox.showerror("Ошибка", "Не удалось обновить клиента!")

    except Exception as e:
        error_msg = str(e)
        if "Duplicate entry" in error_msg:
            messagebox.showerror("Ошибка", "Клиент с таким телефоном уже существует!")
        else:
            messagebox.showerror("Ошибка", f"Не удалось обновить клиента: {error_msg}")


def delete_client_from_db():
    """Удаление выбранного клиента"""
    selection = clients_tree.selection()
    if not selection:
        messagebox.showerror("Ошибка", "Выберите клиента для удаления!")
        return

    item = clients_tree.item(selection[0])
    values = item['values']

    if not values or len(values) < 3:
        messagebox.showerror("Ошибка", "Не удалось получить данные клиента!")
        return

    client_id = values[0]  # Скрытый ID

    # Подтверждение удаления
    if not messagebox.askyesno("Подтверждение", "Вы уверены, что хотите удалить этого клиента?"):
        return

    # Удаляем клиента
    try:
        result = delete_client(client_id)
        if result > 0:
            messagebox.showinfo("Успех", "Клиент удален!")
            refresh_clients_page()

            # Очищаем форму
            client_name_entry.delete(0, tk.END)
            client_phone_entry.delete(0, tk.END)

        else:
            messagebox.showerror("Ошибка", "Не удалось удалить клиента!")

    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось удалить клиента: {str(e)}")


def add_client_from_form():
    """Добавление клиента из формы"""
    name = client_name_entry.get()
    phone = client_phone_entry.get()

    if not name or not phone:
        messagebox.showerror("Ошибка", "Заполните все поля!")
        return

    # Проверяем формат телефона
    if not validate_phone(phone):
        messagebox.showerror("Ошибка", "Неверный формат телефона! Должен быть: + и 11 цифр")
        return

    # Разбиваем имя на части
    name_parts = name.split()
    surname = name_parts[0] if len(name_parts) > 0 else ""
    first_name = name_parts[1] if len(name_parts) > 1 else name_parts[0] if name_parts else ""
    patronymic = name_parts[2] if len(name_parts) > 2 else ""

    # Добавляем клиента
    try:
        add_client(surname, first_name, patronymic, phone)
        messagebox.showinfo("Успех", "Клиент добавлен!")

        # Очищаем форму
        client_name_entry.delete(0, tk.END)
        client_phone_entry.delete(0, tk.END)

        # Обновляем таблицу
        refresh_clients_page()

    except Exception as e:
        error_msg = str(e)
        if "Duplicate entry" in error_msg:
            messagebox.showerror("Ошибка", "Клиент с таким телефоном уже существует!")
        else:
            messagebox.showerror("Ошибка", f"Не удалось добавить клиента: {error_msg}")


def create_clients_page(page):
    """Создание страницы клиентов"""
    page.grid_rowconfigure(2, weight=1)
    page.grid_columnconfigure(0, weight=1)

    # Заголовок
    ttk.Label(page, text="Клиенты", font=LARGE_FONT).grid(row=0, column=0, pady=10, sticky="n")

    # Форма добавления/редактирования
    search_frame = ttk.Frame(page)
    search_frame.grid(row=1, column=0, sticky="ew", padx=12, pady=6)
    search_frame.grid_columnconfigure(0, weight=1)

    ttk.Label(search_frame, text="Имя:", font=MEDIUM_FONT).pack(side="left")
    global client_name_entry
    client_name_entry = ttk.Entry(search_frame, width=25)
    client_name_entry.pack(side="left", padx=6)
    make_large_entry(client_name_entry)

    ttk.Label(search_frame, text="Телефон:", font=MEDIUM_FONT).pack(side="left")
    global client_phone_entry
    client_phone_entry = ttk.Entry(search_frame, width=25)
    client_phone_entry.pack(side="left", padx=6)
    make_large_entry(client_phone_entry)

    # Кнопки формы
    form_buttons_frame = ttk.Frame(search_frame)
    form_buttons_frame.pack(side="left", padx=6)

    ttk.Button(form_buttons_frame, text="Добавить", style="Large.TButton",
               command=add_client_from_form).pack(side="left", padx=2)
    ttk.Button(form_buttons_frame, text="Очистить", style="Large.TButton",
               command=lambda: [client_name_entry.delete(0, tk.END),
                                client_phone_entry.delete(0, tk.END)]).pack(side="left", padx=2)

    # Таблица клиентов
    tree_frame = ttk.Frame(page)
    tree_frame.grid(row=2, column=0, sticky="nsew", padx=12, pady=(6, 0))
    tree_frame.grid_rowconfigure(0, weight=1)
    tree_frame.grid_columnconfigure(0, weight=1)

    # Добавляем столбец для ID (будет скрыт)
    columns = ("id", "name", "phone")
    global clients_tree
    clients_tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=15)
    clients_tree.grid(row=0, column=0, sticky="nsew")

    # Настройка столбцов (первый скрыт)
    clients_tree.heading("name", text="Имя")
    clients_tree.heading("phone", text="Телефон")

    clients_tree.column("name", width=250)
    clients_tree.column("phone", width=150)

    # Привязываем обработчик выбора
    clients_tree.bind("<<TreeviewSelect>>", on_client_select)

    # Scrollbar
    scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=clients_tree.yview)
    scrollbar.grid(row=0, column=1, sticky="ns")
    clients_tree.configure(yscrollcommand=scrollbar.set)

    # Кнопки действий под таблицей
    action_frame = ttk.Frame(page)
    action_frame.grid(row=3, column=0, pady=10, padx=12, sticky="ew")
    action_frame.grid_columnconfigure(0, weight=1)
    action_frame.grid_columnconfigure(1, weight=1)

    ttk.Button(action_frame, text="Изменить", style="Large.TButton",
               command=edit_client).grid(row=0, column=0, padx=4, sticky="ew")
    ttk.Button(action_frame, text="Удалить", style="Large.TButton",
               command=delete_client_from_db).grid(row=0, column=1, padx=4, sticky="ew")

    # Загружаем данные
    refresh_clients_page()


# ========== СТРАНИЦА МАСТЕРОВ ==========
def refresh_masters_page():
    """Обновление страницы мастеров"""
    # Очищаем таблицу
    for item in masters_tree.get_children():
        masters_tree.delete(item)

    # Загружаем мастеров из БД
    masters = get_masters()
    for master in masters:
        # Формируем полное имя
        full_name = f"{master['surname']} {master['name']}"
        if master.get('patronymic'):
            full_name += f" {master['patronymic']}"

        masters_tree.insert("", tk.END, values=(
            master['id'],  # Скрытый ID в первой колонке
            master['surname'],
            master['name'],
            master.get('patronymic', '') or '',
            full_name  # Полное имя для отображения
        ))

    # Скрываем первый (ID) и последний (полное имя) столбцы
    masters_tree.column("#1", width=0, stretch=False)
    masters_tree.heading("#1", text="")
    masters_tree.column("#5", width=0, stretch=False)  # Скрываем полное имя
    masters_tree.heading("#5", text="")


def on_master_select(event):
    """Обработка выбора мастера в таблице"""
    selection = masters_tree.selection()
    if selection:
        item = masters_tree.item(selection[0])
        values = item['values']

        # Заполняем поля формы данными выбранного мастера
        if len(values) >= 4:  # У нас 5 значений, но нам нужны первые 4 (без полного имени)
            master_surname_entry.delete(0, tk.END)
            master_surname_entry.insert(0, values[1])  # Фамилия

            master_name_entry.delete(0, tk.END)
            master_name_entry.insert(0, values[2])  # Имя

            master_patronymic_entry.delete(0, tk.END)
            master_patronymic_entry.insert(0, values[3])  # Отчество


def edit_master():
    """Редактирование выбранного мастера"""
    selection = masters_tree.selection()
    if not selection:
        messagebox.showerror("Ошибка", "Выберите мастера для редактирования!")
        return

    item = masters_tree.item(selection[0])
    values = item['values']
    master_id = values[0]

    # Получаем данные из формы
    surname = master_surname_entry.get()
    name = master_name_entry.get()
    patronymic = master_patronymic_entry.get()

    if not surname or not name:
        messagebox.showerror("Ошибка", "Заполните фамилию и имя!")
        return

    # Обновляем мастера
    try:
        result = update_master(master_id, surname, name, patronymic)
        if result > 0:
            messagebox.showinfo("Успех", "Мастер обновлен!")

            # Обновляем таблицу
            refresh_masters_page()

            # Очищаем форму
            master_surname_entry.delete(0, tk.END)
            master_name_entry.delete(0, tk.END)
            master_patronymic_entry.delete(0, tk.END)

        else:
            messagebox.showerror("Ошибка", "Не удалось обновить мастера!")

    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось обновить мастера: {str(e)}")


def delete_master_from_db():
    """Удаление выбранного мастера"""
    selection = masters_tree.selection()
    if not selection:
        messagebox.showerror("Ошибка", "Выберите мастера для удаления!")
        return

    item = masters_tree.item(selection[0])
    values = item['values']
    master_id = values[0]

    confirm_msg = (
        f"Вы уверены, что хотите удалить мастера {values[1]} {values[2]}?\n\n"
        f"ВНИМАНИЕ: При удалении мастера также удалятся:\n"
        f"• Все его услуги\n"
        f"• Все записи к нему\n"
        f"• Его расписание\n\n"
        f"Это действие нельзя отменить!"
    )

    if not messagebox.askyesno("Подтверждение удаления", confirm_msg):
        return

    try:
        result = delete_master(master_id)
        if result > 0:
            messagebox.showinfo("Успех", "Мастер удален!")

            refresh_masters_page()

            master_surname_entry.delete(0, tk.END)
            master_name_entry.delete(0, tk.END)
            master_patronymic_entry.delete(0, tk.END)

        else:
            messagebox.showerror("Ошибка", "Не удалось удалить мастера!")

    except Exception as e:
        error_msg = str(e)
        if "foreign key constraint" in error_msg.lower():
            messagebox.showerror("Ошибка",
                                 "Невозможно удалить мастера, так как у него есть активные записи или услуги!\n"
                                 "Сначала удалите все записи и услуги этого мастера.")
        else:
            messagebox.showerror("Ошибка", f"Не удалось удалить мастера: {error_msg}")


def add_master_from_form():
    """Добавление мастера из формы"""
    surname = master_surname_entry.get()
    name = master_name_entry.get()
    patronymic = master_patronymic_entry.get()

    if not surname or not name:
        messagebox.showerror("Ошибка", "Заполните фамилию и имя!")
        return

    # Добавляем мастера
    try:
        add_master(surname, name, patronymic)
        messagebox.showinfo("Успех", "Мастер добавлен!")

        # Очищаем форму
        master_surname_entry.delete(0, tk.END)
        master_name_entry.delete(0, tk.END)
        master_patronymic_entry.delete(0, tk.END)

        # Обновляем таблицу на текущей странице
        refresh_masters_page()

    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось добавить мастера: {str(e)}")


def create_masters_page(page):
    """Создание страницы мастеров"""
    page.grid_rowconfigure(2, weight=1)
    page.grid_columnconfigure(0, weight=1)

    # Заголовок
    ttk.Label(page, text="Мастера", font=LARGE_FONT).grid(row=0, column=0, pady=10, sticky="n")

    # Форма добавления/редактирования
    add_frame = ttk.Frame(page, padding=8, relief="groove", borderwidth=1)
    add_frame.grid(row=1, column=0, sticky="nsew", padx=12, pady=6)
    add_frame.grid_columnconfigure(0, weight=0)
    add_frame.grid_columnconfigure(1, weight=1)

    ttk.Label(add_frame, text="Фамилия:", font=MEDIUM_FONT).grid(row=0, column=0, sticky="w", padx=4, pady=4)
    global master_surname_entry
    master_surname_entry = ttk.Entry(add_frame)
    master_surname_entry.grid(row=0, column=1, sticky="ew", padx=4, pady=4)
    make_large_entry(master_surname_entry)

    ttk.Label(add_frame, text="Имя:", font=MEDIUM_FONT).grid(row=1, column=0, sticky="w", padx=4, pady=4)
    global master_name_entry
    master_name_entry = ttk.Entry(add_frame)
    master_name_entry.grid(row=1, column=1, sticky="ew", padx=4, pady=4)
    make_large_entry(master_name_entry)

    ttk.Label(add_frame, text="Отчество:", font=MEDIUM_FONT).grid(row=2, column=0, sticky="w", padx=4, pady=4)
    global master_patronymic_entry
    master_patronymic_entry = ttk.Entry(add_frame)
    master_patronymic_entry.grid(row=2, column=1, sticky="ew", padx=4, pady=4)
    make_large_entry(master_patronymic_entry)

    # Кнопки формы
    form_buttons_frame = ttk.Frame(add_frame)
    form_buttons_frame.grid(row=3, column=0, columnspan=2, pady=6)

    ttk.Button(form_buttons_frame, text="Добавить", style="Large.TButton",
               command=add_master_from_form).pack(side="left", padx=4)
    ttk.Button(form_buttons_frame, text="Очистить", style="Large.TButton",
               command=lambda: [master_surname_entry.delete(0, tk.END),
                                master_name_entry.delete(0, tk.END),
                                master_patronymic_entry.delete(0, tk.END)]).pack(side="left", padx=4)

    # Таблица мастеров
    tree_frame = ttk.Frame(page)
    tree_frame.grid(row=2, column=0, sticky="nsew", padx=12, pady=(6, 0))
    tree_frame.grid_rowconfigure(0, weight=1)
    tree_frame.grid_columnconfigure(0, weight=1)

    # Столбцы таблицы: ID (скрыт), Фамилия, Имя, Отчество, Полное имя (скрыт)
    columns = ("id", "surname", "name", "patronymic", "full_name")
    global masters_tree
    masters_tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=15)
    masters_tree.grid(row=0, column=0, sticky="nsew")

    # Настройка видимых столбцов
    masters_tree.heading("surname", text="Фамилия")
    masters_tree.heading("name", text="Имя")
    masters_tree.heading("patronymic", text="Отчество")

    masters_tree.column("surname", width=150, anchor="w")
    masters_tree.column("name", width=150, anchor="w")
    masters_tree.column("patronymic", width=150, anchor="w")

    # Привязываем обработчик выбора
    masters_tree.bind("<<TreeviewSelect>>", on_master_select)

    # Scrollbar
    scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=masters_tree.yview)
    scrollbar.grid(row=0, column=1, sticky="ns")
    masters_tree.configure(yscrollcommand=scrollbar.set)

    # Кнопки действий под таблицей
    action_frame = ttk.Frame(page)
    action_frame.grid(row=3, column=0, pady=10, padx=12, sticky="ew")
    action_frame.grid_columnconfigure(0, weight=1)
    action_frame.grid_columnconfigure(1, weight=1)
    action_frame.grid_columnconfigure(2, weight=1)
    action_frame.grid_columnconfigure(3, weight=1)

    ttk.Button(action_frame, text="Изменить", style="Large.TButton",
               command=edit_master).grid(row=0, column=0, padx=4, sticky="ew")
    ttk.Button(action_frame, text="Удалить", style="Large.TButton",
               command=delete_master_from_db).grid(row=0, column=1, padx=4, sticky="ew")
    ttk.Button(action_frame, text="Назначить услугу", style="Large.TButton",
               command=set_service_to_master).grid(row=0, column=2, padx=4, sticky="ew")
    ttk.Button(action_frame, text="Освободить от услуги", style="Large.TButton",
               command=unset_service_to_master).grid(row=0, column=3, padx=4, sticky="ew")

    # Загружаем данные
    refresh_masters_page()


def set_service_to_master():
    """Назначает мастеру услугу"""
    def set_service():
        nonlocal master_id, avaliable_service_combo, set_service_window
        service_id = avaliable_service_combo.get().split("|")[0]
        set_master_service_to_db(master_id, service_id)
        set_service_window.destroy()

    try:
        selection = masters_tree.selection()
        if not selection:
            messagebox.showerror("Ошибка", "Выберите мастера для назначения услуги!")
            return 0

        item = masters_tree.item(selection[0])
        values = item['values']
        master_id = values[0]
        master_full_name = values[4]

        unsetted_master_service = get_unsetted_master_service(master_id)
        if len(unsetted_master_service) == 0:
            messagebox.showwarning("Внимание", "Данному мастеру назначены все существующие услуги!")
            return 0

        avaliable_service_combo_values = [f"{service_record["id"]}|{service_record["name"]}|{service_record["duration"]} мин." for service_record in unsetted_master_service]

        set_service_window = tk.Toplevel(master=root)
        set_service_window.title("Назначить услугу мастеру")

        avaliable_service_combo = ttk.Combobox(
            master=set_service_window,
            values=avaliable_service_combo_values,
            state="readonly",
            font=SMALL_FONT,
        )

        tk.Label(
            master=set_service_window,
            text=f"Назначить услугу для мастера\n\"{master_full_name}\"",
            font=MEDIUM_FONT
        ).grid(row=0, column=0, sticky="we", padx=10, pady=2)

        avaliable_service_combo.grid(row=1, column=0, sticky="we", padx=10, pady=2)

        ttk.Button(
            master=set_service_window,
            text="Назначить услугу",
            style="Small.TButton",
            command=set_service
        ).grid(row=2, column=0, sticky="we", padx=10, pady=2)

        set_service_window.grab_set()
        set_service_window.wait_window()

    except Exception as e:
        messagebox.showerror("Ошибка", "Произошла ошибка при подключении к БД.")


def unset_service_to_master():
    """Освобождает мастера от услуги"""
    def unset_service():
        nonlocal master_id, set_service_combo, unset_service_window
        service_id = set_service_combo.get().split("|")[0]
        unset_master_service_to_db(master_id, service_id)
        unset_service_window.destroy()

    try:
        selection = masters_tree.selection()
        if not selection:
            messagebox.showerror("Ошибка", "Выберите мастера для отмены услуги!")
            return 0

        item = masters_tree.item(selection[0])
        values = item['values']
        master_id = values[0]
        master_full_name = values[4]

        set_master_service = get_setted_master_service(master_id)

        if len(set_master_service) == 0:
            messagebox.showwarning("Внимание", "Данному мастеру не назначено ни одной услуги!")
            return 0


        set_service_combo_values = [f"{service_record["id"]}|{service_record["name"]}|{service_record["duration"]} мин." for service_record in set_master_service]

        unset_service_window = tk.Toplevel(master=root)
        unset_service_window.title("Освободить мастера от услуги")

        set_service_combo = ttk.Combobox(
            master=unset_service_window,
            values=set_service_combo_values,
            state="readonly",
            font=SMALL_FONT,
        )

        tk.Label(
            master=unset_service_window,
            text=f"Освободить от услуги мастера\n\"{master_full_name}\"",
            font=MEDIUM_FONT
        ).grid(row=0, column=0, sticky="we", padx=10, pady=2)

        set_service_combo.grid(row=1, column=0, sticky="we", padx=10, pady=2)

        ttk.Button(
            master=unset_service_window,
            text="Отменить услугу",
            style="Small.TButton",
            command=unset_service
        ).grid(row=2, column=0, sticky="we", padx=10, pady=2)

        unset_service_window.grab_set()
        unset_service_window.wait_window()
    except Exception as e:
        messagebox.showerror("Ошибка", "Произошла ошибка при подключении к БД.")


#------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------
#                       ОТЧЕТЫ

# Возвращает границы периода для отчета
def period_parse(period_string_value: str) -> tuple:
    current_date = date.today()
    match period_string_value:
        case "day":
            start_date = current_date
            end_date = current_date
            return start_date, end_date
        case "month":
            current_year, current_month = current_date.year, current_date.month

            if current_month != 12:
                end_month, end_year = current_month + 1, current_year
            else:
                end_month, end_year = 1, current_year + 1
            start_date = datetime.strptime(f"{current_year}-{current_month}-01", "%Y-%m-%d")
            end_date = datetime.strptime(f"{end_year}-{end_month}-01", "%Y-%m-%d") - timedelta(days=1)
            return start_date, end_date
        case "year":
            current_year = current_date.year
            end_year = current_year + 1

            start_date = datetime.strptime(f"{current_year}-01-01", "%Y-%m-%d")
            end_date = datetime.strptime(f"{end_year}-01-01", "%Y-%m-%d") - timedelta(days=1)
            return start_date, end_date


def generate_attendance_report(start_report_date, end_report_date):
    """Генерация отчета по посещаемости (оптимизированная SQL версия)"""
    try:
        # Используем SQL-запрос вместо Python-обработки
        stats = get_attendance_statistics(start_report_date, end_report_date)

        if not stats or stats['total'] == 0:
            return "Нет данных для отчета"

        total = stats['total'] or 0
        completed = stats['completed'] or 0
        cancelled = stats['cancelled'] or 0
        pending = stats['pending'] or 0
        in_progress = stats['in_progress'] or 0

        attendance_rate = (completed / total * 100) if total > 0 else 0
        cancellation_rate = (cancelled / total * 100) if total > 0 else 0

        report = f"""
    ОТЧЕТ ПО ПОСЕЩАЕМОСТИ ЗА ПЕРИОД
    ================================
    Период: {start_report_date.strftime('%d.%m.%Y')} - {end_report_date.strftime('%d.%m.%Y')}

    Общее количество записей: {total}
    Выполнено: {completed} ({attendance_rate:.1f}%)
    Отменено: {cancelled} ({cancellation_rate:.1f}%)
    Ожидается: {pending}
    В процессе: {in_progress}

    Анализ:
    - Уровень посещаемости: {attendance_rate:.1f}%
    - Процент отмен: {cancellation_rate:.1f}%
    - Эффективность работы: {(completed / (total - cancelled) * 100 if total - cancelled > 0 else 0):.1f}%
    """

        return report
    except Exception as e:
        return f"Ошибка при генерации отчета: {str(e)}"


def generate_master_profit_report(start_report_date, end_report_date):
    """Генерация отчета по прибыли по мастерам (оптимизированная SQL версия)"""
    try:
        # Используем SQL-запрос вместо Python-обработки
        master_profits = get_master_profit_report(start_report_date, end_report_date)

        if not master_profits:
            return "Нет данных для отчета"

        report = "ОТЧЕТ ПО ПРИБЫЛИ ПО МАСТЕРАМ\n"
        report += "============================\n"
        report += f"Период: {start_report_date.strftime('%d.%m.%Y')} - {end_report_date.strftime('%d.%m.%Y')}\n\n"

        for master_data in master_profits:
            report += f"{master_data['master_name']}:\n"
            report += f"  Количество выполненных услуг: {master_data['service_count']}\n"
            report += f"  Общая прибыль: {master_data['total_profit'] or 0:.2f} руб.\n"
            report += f"  Средняя прибыль за услугу: {master_data['avg_profit'] or 0:.2f} руб.\n"

            # Рассчитываем долю в общей прибыли
            total_all = sum(m['total_profit'] or 0 for m in master_profits)
            if total_all > 0:
                percentage = (master_data['total_profit'] or 0) / total_all * 100
                report += f"  Доля в общей прибыли: {percentage:.1f}%\n"

            report += "\n"

        # Добавляем итоговую статистику
        total_services = sum(m['service_count'] or 0 for m in master_profits)
        total_profit = sum(m['total_profit'] or 0 for m in master_profits)
        avg_profit_all = total_profit / total_services if total_services > 0 else 0

        report += f"ИТОГО:\n"
        report += f"  Общее количество услуг: {total_services}\n"
        report += f"  Общая прибыль: {total_profit:.2f} руб.\n"
        report += f"  Средняя прибыль за услугу: {avg_profit_all:.2f} руб.\n"

        return report
    except Exception as e:
        return f"Ошибка при генерации отчета: {str(e)}"


def generate_top_services_report(start_report_date, end_report_date):
    """Генерация отчета по топ-5 услуг (оптимизированная SQL версия)"""
    try:
        # Используем SQL-запрос вместо Python-обработки
        top_services = get_top_services_report(start_report_date, end_report_date)

        if not top_services:
            return "Нет данных для отчета"

        report = "ТОП-5 САМЫХ ПОПУЛЯРНЫХ УСЛУГ\n"
        report += "============================\n"
        report += f"Период: {start_report_date.strftime('%d.%m.%Y')} - {end_report_date.strftime('%d.%m.%Y')}\n\n"

        # Получаем общее количество выполненных услуг за период
        stats = get_attendance_statistics(start_report_date, end_report_date)
        total_completed = stats['completed'] or 0 if stats else 0

        for i, service_data in enumerate(top_services, 1):
            service_name = service_data['service_name']
            count = service_data['service_count'] or 0

            # Рассчитываем процент от общего количества
            percentage = (count / total_completed * 100) if total_completed > 0 else 0

            report += f"{i}. {service_name}\n"
            report += f"   Количество выполненных: {count}\n"
            report += f"   Доля от всех услуг: {percentage:.1f}%\n\n"

        if total_completed > 0:
            report += f"Всего выполнено услуг за период: {total_completed}\n"

        return report
    except Exception as e:
        return f"Ошибка при генерации отчета: {str(e)}"


def generate_master_load_report(start_report_date, end_report_date):
    """Генерация отчета по загруженности мастеров (оптимизированная SQL версия)"""
    try:
        # Используем SQL-запрос вместо Python-обработки
        master_loads = get_master_load_report(start_report_date, end_report_date)

        if not master_loads:
            return "Нет данных для отчета"

        report = "ОТЧЕТ ПО ЗАГРУЖЕННОСТИ МАСТЕРОВ\n"
        report += "===============================\n"
        report += f"Период: {start_report_date.strftime('%d.%m.%Y')} - {end_report_date.strftime('%d.%m.%Y')}\n\n"

        # Группируем данные по мастерам
        masters_data = {}
        for load_data in master_loads:
            master_name = load_data['master_name']
            status = load_data['status']
            count = load_data['count'] or 0

            if master_name not in masters_data:
                masters_data[master_name] = {'total': 0, 'by_status': {}}

            masters_data[master_name]['by_status'][status] = count
            masters_data[master_name]['total'] += count

        for master_name, data in masters_data.items():
            report += f"{master_name}:\n"
            report += f"  Всего записей: {data['total']}\n"

            # Выводим по статусам
            for status, count in data['by_status'].items():
                percentage = (count / data['total'] * 100) if data['total'] > 0 else 0
                report += f"  {status}: {count} ({percentage:.1f}%)\n"

            # Рассчитываем загруженность (отношение выполненных к общему числу)
            completed = data['by_status'].get('Выполнено', 0)
            effective_load = (completed / data['total'] * 100) if data['total'] > 0 else 0
            report += f"  Эффективная загруженность: {effective_load:.1f}%\n\n"

        return report
    except Exception as e:
        return f"Ошибка при генерации отчета: {str(e)}"


def generate_appointments_by_master_report(start_report_date, end_report_date):
    """Генерация отчета по количеству записей по мастерам (оптимизированная SQL версия)"""
    try:
        # Используем SQL-запрос вместо Python-обработки
        appointments_data = get_appointments_by_master_detailed(start_report_date, end_report_date)

        if not appointments_data:
            return "Нет данных для отчета"

        report = "КОЛИЧЕСТВО ЗАПИСЕЙ ПО МАСТЕРАМ ЗА ПЕРИОД\n"
        report += "======================================\n"
        report += f"Период: {start_report_date.strftime('%d.%m.%Y')} - {end_report_date.strftime('%d.%m.%Y')}\n\n"

        # Группируем данные по мастерам
        masters_summary = {}
        for app_data in appointments_data:
            master_name = app_data['master_name']
            appointment_day = app_data['appointment_day']
            count = app_data['appointment_count'] or 0

            if master_name not in masters_summary:
                masters_summary[master_name] = {'total': 0, 'by_day': {}}

            masters_summary[master_name]['by_day'][appointment_day] = count
            masters_summary[master_name]['total'] += count

        # Сортируем мастеров по общему количеству записей
        sorted_masters = sorted(masters_summary.items(), key=lambda x: x[1]['total'], reverse=True)

        for master_name, data in sorted_masters:
            report += f"{master_name}:\n"

            # Сортируем дни по дате
            sorted_days = sorted(data['by_day'].items(), key=lambda x: x[0])

            for day, count in sorted_days:
                report += f"  {day.strftime('%d.%m.%Y')}: {count} записей\n"

            # Рассчитываем среднее количество записей в день
            days_count = len(data['by_day'])
            avg_per_day = data['total'] / days_count if days_count > 0 else 0

            report += f"  Итого: {data['total']} записей за {days_count} дней\n"
            report += f"  В среднем в день: {avg_per_day:.1f} записей\n\n"

        # Общая статистика
        total_appointments = sum(data['total'] for _, data in sorted_masters)
        total_masters = len(sorted_masters)
        avg_per_master = total_appointments / total_masters if total_masters > 0 else 0

        report += f"ОБЩАЯ СТАТИСТИКА:\n"
        report += f"  Всего мастеров: {total_masters}\n"
        report += f"  Всего записей: {total_appointments}\n"
        report += f"  В среднем на мастера: {avg_per_master:.1f} записей\n"

        return report
    except Exception as e:
        return f"Ошибка при генерации отчета: {str(e)}"


def save_report_to_docx(report_text, report_type):
    """Сохранение отчета в файл .docx"""
    # Открываем диалог выбора файла
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
        initialfile=f"{report_type}.docx"
    )

    if not file_path:
        return False  # Пользователь отменил сохранение

    try:
        # Создаем документ
        doc = Document()

        # Добавляем заголовок
        title = doc.add_heading(report_type, 0)
        title.alignment = 1  # Центрирование

        # Добавляем дату
        from datetime import datetime
        date_para = doc.add_paragraph()
        date_para.add_run(f"Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}").italic = True
        date_para.alignment = 1

        doc.add_paragraph()  # Пустая строка

        # Разбиваем текст на параграфы и добавляем в документ
        lines = report_text.split('\n')
        for line in lines:
            if line.strip() == '':
                doc.add_paragraph()  # Пустая строка
            elif line.startswith('='):
                # Заголовок второго уровня
                heading = doc.add_heading(line.strip('= '), level=2)
            elif ':' in line and not line.startswith('  '):
                # Важные пункты
                p = doc.add_paragraph()
                p.add_run(line).bold = True
            else:
                # Обычный текст
                doc.add_paragraph(line)

        # Сохраняем документ
        doc.save(file_path)

        # Открываем файл (опционально)
        if os.name == 'nt':  # Windows
            os.startfile(file_path)
        elif os.name == 'posix':  # Linux, Mac
            os.system(f'open "{file_path}"' if sys.platform == 'darwin' else f'xdg-open "{file_path}"')

        return True

    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось сохранить отчет: {str(e)}")
        return False


def show_report_dialog(report_type, report_function):
    """Показать диалоговое окно с отчетом"""
    # Создаем новое окно
    report_window = tk.Toplevel(root)
    report_window.title(f"Отчет: {report_type}")
    report_window.geometry("800x600")
    report_window.transient(root)  # Сделать окно дочерним

    # Создаем текстовое поле для отчета
    text_frame = ttk.Frame(report_window)
    text_frame.pack(fill="both", expand=True, padx=10, pady=10)

    text_widget = tk.Text(text_frame, wrap="word", font=("Courier", 10))
    scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=text_widget.yview)
    text_widget.configure(yscrollcommand=scrollbar.set)

    text_widget.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # Получаем выбранный период
    global selected_report_period
    current_period = selected_report_period.get()
    start_report_date, end_report_date = period_parse(current_period)

    # Генерируем отчет
    report_text = report_function(start_report_date, end_report_date)
    text_widget.insert("1.0", report_text)
    text_widget.configure(state="disabled")  # Запрещаем редактирование

    # Фрейм с кнопками
    button_frame = ttk.Frame(report_window)
    button_frame.pack(fill="x", padx=10, pady=(0, 10))

    # Кнопка сохранения
    def save_report():
        if save_report_to_docx(report_text, report_type):
            messagebox.showinfo("Успех", f"Отчет сохранен!")

    ttk.Button(button_frame, text="Сохранить в Word",
               command=save_report, style="Large.TButton").pack(side="right", padx=5)

    # Кнопка закрытия
    ttk.Button(button_frame, text="Закрыть",
               command=report_window.destroy).pack(side="right", padx=5)

# ========== СТРАНИЦА ОТЧЕТОВ ==========
def create_reports_page(page):
    """Создание страницы отчетов"""
    page.grid_rowconfigure(0, weight=1)
    page.grid_columnconfigure(0, weight=1)

    report_frame = ttk.Frame(page, padding=12)
    report_frame.grid(row=0, column=0, sticky="nsew")

    report_frame.grid_rowconfigure(4, weight=1)
    report_frame.grid_columnconfigure(index=1, weight=1)

    ttk.Label(report_frame, text="Отчеты", font=LARGE_FONT).grid(
        row=0, column=0, columnspan=3, pady=(0, 10), sticky="n"
    )

    ttk.Label(report_frame, text="Выберите период", font=MEDIUM_FONT).grid(
        row=1, column=0, sticky="w", padx=(0, 8)
    )

    global selected_report_period
    selected_report_period = tk.StringVar(value="day")

    rad_frame = ttk.Frame(report_frame)
    rad_frame.grid(row=1, column=1, columnspan=2, sticky="w")
    ttk.Radiobutton(rad_frame, text="день", value="day", variable=selected_report_period).pack(side="left", padx=6)
    ttk.Radiobutton(rad_frame, text="месяц", value="month", variable=selected_report_period).pack(side="left", padx=6)
    ttk.Radiobutton(rad_frame, text="год", value="year", variable=selected_report_period).pack(side="left", padx=6)

    ttk.Separator(report_frame, orient="horizontal").grid(
        row=2, column=0, columnspan=2, sticky="ew", pady=10
    )

    ttk.Label(report_frame, text="Тип отчета", font=MEDIUM_FONT).grid(
        row=3, column=0, columnspan=2, sticky="w", pady=(0, 6)
    )

    report_types = [
        ("Процент посещаемости", generate_attendance_report),
        ("Прибыль по мастерам", generate_master_profit_report),
        ("Топ 5 самых популярных услуг", generate_top_services_report),
        ("Загруженность мастеров", generate_master_load_report),
        ("Количество записей по мастерам", generate_appointments_by_master_report)
    ]

    report_buttons_frame = ttk.Frame(report_frame)
    report_buttons_frame.grid(row=4, column=0, columnspan=2, sticky="nsew", pady=(0, 8))
    report_buttons_frame.grid_columnconfigure(index=1, weight=1)

    for i, (text, func) in enumerate(report_types):
        # Метка с названием отчета
        ttk.Label(
            report_buttons_frame,
            text=text, font=SMALL_FONT
        ).grid(row=i, column=0, sticky="w", pady=6, padx=(0, 10))

        # Кнопка для формирования отчета
        ttk.Button(
            report_buttons_frame,
            text="Сформировать",
            style="Large.TButton",
            command=lambda f=func, t=text: show_report_dialog(t, f)
        ).grid(row=i, column=1, sticky="w", pady=6, ipady=8)


# ========== ОСНОВНОЙ ИНТЕРФЕЙС ==========
def main():
    """Основная функция приложения"""
    global root

    # Создаем главный контейнер
    main_container = ttk.Frame(root)
    main_container.pack(fill="both", expand=True)

    # Левое меню
    menu_frame = ttk.Frame(main_container, width=260)
    menu_frame.pack(side="left", fill="y")
    menu_frame.pack_propagate(False)

    # Область контента
    content_frame = ttk.Frame(main_container)
    content_frame.pack(side="right", fill="both", expand=True)

    # Создаем кнопки меню
    for name in page_names:
        btn = ttk.Button(menu_frame, text=name, style="Menu.TButton",
                         command=lambda n=name: show_page(n))
        btn.pack(fill="x", pady=4, ipady=8, expand=True)
        menu_buttons[name] = btn

    # Кнопка выхода
    ttk.Button(menu_frame, text="Выход", style="Menu.TButton",
               command=quit_app).pack(fill="x", pady=(8, 8), ipady=8)

    # Создаем страницы
    for name in page_names:
        frame = ttk.Frame(content_frame)
        frame.grid(row=0, column=0, sticky="nsew")
        pages[name] = frame

    content_frame.grid_rowconfigure(0, weight=1)
    content_frame.grid_columnconfigure(0, weight=1)

    # Настраиваем стиль таблиц
    style.configure("Treeview", rowheight=40)

    # Создаем все страницы
    create_main_page(pages["Главная"])
    create_booking_page(pages["Запись"])
    create_schedule_page(pages["График работы"])
    create_services_page(pages["Услуги"])
    create_clients_page(pages["Клиенты"])
    create_masters_page(pages["Мастера"])
    create_reports_page(pages["Отчеты"])

    # Показываем главную страницу
    show_page("Главная")

    # Запускаем главный цикл
    root.mainloop()



if __name__ == "__main__":
    main()